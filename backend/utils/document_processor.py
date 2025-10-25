"""
Document Processor Module
Extracts text from PDF, DOCX, and TXT files
"""

import fitz  # PyMuPDF
from PyPDF2 import PdfReader
from docx import Document
import warnings
import re
import os

# Suppress noisy MuPDF warnings
warnings.filterwarnings("ignore", category=UserWarning, module="fitz")


class DocumentProcessor:
    """Handles text extraction from various document formats"""

    def __init__(self):
        """Initialize the document processor"""
        self.supported_formats = ['pdf', 'docx', 'txt']

    def process_document(self, file_path):
        """
        Main method to process a document and extract text

        Args:
            file_path (str): Path to the document file

        Returns:
            dict: Dictionary containing extracted text and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = file_path.lower().split('.')[-1]

        if file_extension == 'pdf':
            return self._process_pdf(file_path)
        elif file_extension == 'docx':
            return self._process_docx(file_path)
        elif file_extension == 'txt':
            return self._process_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    # ----------------------------------------------------------------------
    # PDF PROCESSING (Robust with fallback)
    # ----------------------------------------------------------------------
    def _process_pdf(self, file_path):
        """
        Extract text from PDF file using PyMuPDF with PyPDF2 fallback
        """
        text = ""
        page_count = 0

        # --- Try PyMuPDF first ---
        try:
            doc = fitz.open(file_path)
            page_count = len(doc)

            for page in doc:
                text += page.get_text("text")
            doc.close()

            if text.strip():
                cleaned_text = self._clean_text(text)
                return {
                    'text': cleaned_text,
                    'page_count': page_count,
                    'word_count': len(cleaned_text.split()),
                    'char_count': len(cleaned_text),
                    'format': 'pdf'
                }
            else:
                print("⚠️ PyMuPDF extracted no text, trying PyPDF2...")
        except Exception as e:
            print(f"⚠️ MuPDF failed: {e}. Trying PyPDF2...")

        # --- Fallback: PyPDF2 ---
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            page_count = len(reader.pages)

            if text.strip():
                cleaned_text = self._clean_text(text)
                print("✅ Successfully extracted text using PyPDF2 fallback.")
                return {
                    'text': cleaned_text,
                    'page_count': page_count,
                    'word_count': len(cleaned_text.split()),
                    'char_count': len(cleaned_text),
                    'format': 'pdf'
                }
        except Exception as e2:
            print(f"❌ PDF extraction failed completely: {e2}")

        # --- No text extracted ---
        return {
            'text': "",
            'page_count': page_count,
            'word_count': 0,
            'char_count': 0,
            'format': 'pdf'
        }

    # ----------------------------------------------------------------------
    # DOCX PROCESSING
    # ----------------------------------------------------------------------
    def _process_docx(self, file_path):
        """
        Extract text from DOCX file using python-docx
        """
        try:
            doc = Document(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            cleaned_text = self._clean_text(text)

            return {
                'text': cleaned_text,
                'paragraph_count': len(doc.paragraphs),
                'word_count': len(cleaned_text.split()),
                'char_count': len(cleaned_text),
                'format': 'docx'
            }

        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")

    # ----------------------------------------------------------------------
    # TXT PROCESSING
    # ----------------------------------------------------------------------
    def _process_txt(self, file_path):
        """
        Extract text from TXT file with encoding fallback
        """
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            text = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                raise Exception("Could not decode file with any supported encoding")

            cleaned_text = self._clean_text(text)
            line_count = len([line for line in text.split('\n') if line.strip()])

            return {
                'text': cleaned_text,
                'line_count': line_count,
                'word_count': len(cleaned_text.split()),
                'char_count': len(cleaned_text),
                'format': 'txt'
            }

        except Exception as e:
            raise Exception(f"Error processing TXT: {str(e)}")

    # ----------------------------------------------------------------------
    # TEXT CLEANING
    # ----------------------------------------------------------------------
    def _clean_text(self, text):
        """
        Clean and normalize extracted text
        """
        if not text:
            return ""

        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\'\"]', '', text)
        text = re.sub(r'\.{2,}', '.', text)
        text = re.sub(r' {2,}', ' ', text)
        text = text.strip()
        return text

    # ----------------------------------------------------------------------
    # VALIDATION & PREVIEW
    # ----------------------------------------------------------------------
    def get_text_preview(self, text, max_length=200):
        """
        Get a short preview of the text
        """
        if len(text) <= max_length:
            return text

        preview = text[:max_length]
        last_period = preview.rfind('.')

        if last_period > max_length * 0.7:
            preview = preview[:last_period + 1]
        else:
            preview += "..."
        return preview

    def validate_text(self, text, min_words=50):
        """
        Validate extracted text for minimum quality
        """
        if not text or len(text.strip()) == 0:
            return False, "No text could be extracted from the document"

        word_count = len(text.split())
        if word_count < min_words:
            return False, f"Document too short: {word_count} words (min {min_words})"

        return True, "Text is valid"


# ----------------------------------------------------------------------
# UTILITY FUNCTION
# ----------------------------------------------------------------------
def extract_text(file_path):
    """
    Convenience function to extract text from any supported document
    """
    processor = DocumentProcessor()
    return processor.process_document(file_path)
