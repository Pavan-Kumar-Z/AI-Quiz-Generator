from docx import Document

# Create a new Document
doc = Document()

# Add title
doc.add_heading('Test Document for Quiz Generator', 0)

# Add paragraphs
doc.add_paragraph('This is a test document for the AI Quiz Generator application.')
doc.add_paragraph('Python is a versatile programming language used in many fields.')

doc.add_heading('Key Features of Python', level=1)
doc.add_paragraph('1. Easy to learn and read')
doc.add_paragraph('2. Large standard library')
doc.add_paragraph('3. Cross-platform compatibility')
doc.add_paragraph('4. Strong community support')

doc.add_heading('Applications', level=1)
doc.add_paragraph('Python is used in web development, data science, machine learning, automation, and more.')

# Save
doc.save('test_document.docx')
print("✅ Created test_document.docx")